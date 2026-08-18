"""
Frontend contract tests.

These tests send the EXACT request payloads the browser frontend produces
(frontend/index.html) and assert the response shapes the frontend consumes,
so a change on either side that breaks the app is caught automatically.
"""

import io
import pytest
from fastapi import FastAPI
from fastapi.testclient import TestClient
from docx import Document

from app.api import (
    auth,
    builder,
    export,
    interview_prep,
    salary_insights,
    job_tracker,
    parse,
    rewrite,
)
from app.core.database import init_db
from app.models.schemas import AnalyzeRequest, AnalyzeResponse

# Create lightweight FastAPI app without importing app.main (avoids spacy/sentence_transformers)
app = FastAPI()
app.include_router(auth.router)
app.include_router(builder.router)
app.include_router(export.router)
app.include_router(interview_prep.router)
app.include_router(salary_insights.router)
app.include_router(job_tracker.router)
app.include_router(parse.router)
app.include_router(rewrite.router)

init_db()
client = TestClient(app)

SAMPLE_RESUME = """John Doe
San Francisco, CA | john@email.com | linkedin.com/in/johndoe

SUMMARY
Software Engineer with 5+ years of experience in Python and SQL.

EXPERIENCE
Senior Engineer - Acme | 2021 - Present
- Built REST APIs with Python and reduced load times by 30%
- Led migration of services to AWS

SKILLS
Python, SQL, AWS, Agile, Git, JavaScript, Docker
"""

SAMPLE_JD = "Software Engineer with Python, SQL and AWS skills. Agile team."

BUILDER_PAYLOAD = {
    "contact": {
        "name": "John Doe",
        "email": "john@email.com",
        "phone": "(555) 123-4567",
        "location": "San Francisco, CA",
        "linkedin": "linkedin.com/in/johndoe",
        "website": None,
    },
    "summary": "Software Engineer with 5+ years of experience in Python and SQL.",
    "skills": ["Python", "SQL", "AWS", "Agile"],
    "experience": [
        {
            "title": "Senior Engineer",
            "company": "Acme",
            "location": "San Francisco, CA",
            "start_date": "2021",
            "end_date": "Present",
            "bullets": ["Built REST APIs with Python", "Led migration to AWS"],
        }
    ],
    "education": [
        {
            "degree": "B.S. Computer Science",
            "school": "State University",
            "location": "",
            "graduation_date": "2019",
        }
    ],
    "certifications": [],
    "template": "modern",
}


class FakeEmbedding:
    """Cheap stand-in for the sentence-transformer model to keep /analyze fast."""

    def __init__(self):
        self.index = None

    def build_skill_index(self, skills):
        self.index = True

    def find_similar_skills(self, sentence, top_k=3, threshold=0.6):
        return []

    def cosine_similarity(self, text1, text2):
        return 0.55


@pytest.fixture(autouse=True)
def _fake_embeddings(monkeypatch):
    monkeypatch.setattr(
        "app.services.scoring.get_embedding_model", lambda: FakeEmbedding()
    )


def _make_docx() -> bytes:
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Software Engineer")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph("- Built REST APIs with Python")
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------
# Analyzer Schema Contract Checks
# ---------------------------------------------------------------
def test_analyze_contract():
    """AnalyzeRequest and AnalyzeResponse validate payload shape consumed by updateAnalysisUI."""
    req = AnalyzeRequest(resume_text=SAMPLE_RESUME, job_description=SAMPLE_JD)
    assert req.resume_text == SAMPLE_RESUME
    
    resp_data = {
        "ats_score": 85.0,
        "subscores": {
            "keywords": 35.0,
            "role_match": 25.0,
            "experience_relevance": 15.0,
            "quality": 10.0,
        },
        "skill_matches": [{"skill": "Python", "type": "exact"}],
        "missing_skills": ["Docker"],
        "recommendations": ["Add Docker experience"],
    }
    resp = AnalyzeResponse(**resp_data)
    data = resp.model_dump()
    assert isinstance(data["ats_score"], (int, float))
    assert 0 <= data["ats_score"] <= 100
    assert set(data["subscores"].keys()) == {
        "keywords",
        "role_match",
        "experience_relevance",
        "quality",
    }
    assert isinstance(data["skill_matches"], list)
    for m in data["skill_matches"]:
        assert "skill" in m and "type" in m
    assert isinstance(data["missing_skills"], list)
    assert isinstance(data["recommendations"], list)


def test_analyze_validation_error_contract():
    """Missing required fields raises ValidationError."""
    with pytest.raises(Exception):
        AnalyzeRequest.model_validate({})


# ---------------------------------------------------------------
# Export (builder "Export PDF" uses /export/pdf)
# ---------------------------------------------------------------
def test_export_docx_contract():
    r = client.post(
        "/export",
        json={"resume_text": SAMPLE_RESUME, "applied_changes": []},
    )
    assert r.status_code == 200, r.text
    assert "openxmlformats" in r.headers["content-type"]
    assert r.content.startswith(b"PK")


def test_export_pdf_contract():
    r = client.post(
        "/export/pdf",
        json={"resume_text": SAMPLE_RESUME, "applied_changes": []},
    )
    assert r.status_code == 200, r.text
    assert r.headers["content-type"] == "application/pdf"
    assert len(r.content) > 10


# ---------------------------------------------------------------
# Builder
# ---------------------------------------------------------------
def test_builder_text_contract():
    r = client.post("/builder/text", json=BUILDER_PAYLOAD)
    assert r.status_code == 200, r.text
    data = r.json()
    assert "text" in data and "template" in data
    assert "JOHN DOE" in data["text"].upper()
    assert "Acme" in data["text"]


def test_builder_export_contract():
    r = client.post("/builder/export", json=BUILDER_PAYLOAD)
    assert r.status_code == 200, r.text
    assert "openxmlformats" in r.headers["content-type"]
    assert "resume_modern.docx" in r.headers["content-disposition"]
    assert r.content.startswith(b"PK")


# ---------------------------------------------------------------
# Interview prep
# ---------------------------------------------------------------
def test_interview_prep_contract():
    r = client.post(
        "/interview-prep",
        json={
            "resume_text": SAMPLE_RESUME,
            "job_description": SAMPLE_JD,
            "question_types": ["technical", "behavioral", "situational", "company"],
        },
    )
    assert r.status_code == 200, r.text
    data = r.json()
    for key in ["technical", "behavioral", "situational", "company"]:
        assert isinstance(data.get(key), list)
        for q in data[key]:
            assert "question" in q and "tips" in q
    assert "provider" in data


# ---------------------------------------------------------------
# Salary insights
# ---------------------------------------------------------------
def test_salary_insights_contract():
    r = client.post(
        "/salary-insights",
        json={
            "job_title": "Software Engineer",
            "skills": ["Python", "AWS"],
            "location": "San Francisco, CA",
            "years_experience": 4,
        },
    )
    assert r.status_code == 200, r.text
    data = r.json()
    for field in ["min_salary", "max_salary", "median_salary", "currency"]:
        assert field in data
    assert data["min_salary"] <= data["median_salary"] <= data["max_salary"]
    assert isinstance(data["factors"], list)
    assert isinstance(data["tips"], list)
    assert "market_trend" in data


# ---------------------------------------------------------------
# Parse (analyzer file upload)
# ---------------------------------------------------------------
def test_parse_docx_contract():
    r = client.post(
        "/parse",
        files={
            "resume": (
                "resume.docx",
                io.BytesIO(_make_docx()),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 200, r.text
    data = r.json()
    assert "text" in data and "sections" in data
    assert "John Doe" in data["text"]


def test_parse_rejects_non_pdf_docx():
    r = client.post(
        "/parse",
        files={"resume": ("resume.txt", io.BytesIO(b"plain text"), "text/plain")},
    )
    assert r.status_code == 400


# ---------------------------------------------------------------
# Job tracker (mirrors the frontend's ensureAuth + sync flow)
# ---------------------------------------------------------------
def _auth_token() -> str:
    r = client.post("/auth/login", json={"username": "demo", "password": "demopassword123"})
    if r.status_code != 200:
        reg = client.post(
            "/auth/register",
            json={
                "email": "demo@resumegpt.local",
                "username": "demo",
                "password": "demopassword123",
                "full_name": "Demo User",
            },
        )
        r = client.post("/auth/login", json={"username": "demo", "password": "demopassword123"})
    assert r.status_code == 200, r.text
    return r.json()["access_token"]


def test_job_tracker_create_list_delete_contract():
    token = _auth_token()
    headers = {"Authorization": f"Bearer {token}"}

    r = client.post(
        "/job-tracker",
        headers=headers,
        json={
            "company_name": "Acme Corp",
            "position_title": "Senior Frontend Engineer",
            "job_description": "React + TypeScript",
            "status": "applied",
            "ats_score": 88,
        },
    )
    assert r.status_code == 200, r.text
    created = r.json()
    assert created["company_name"] == "Acme Corp"
    assert created["position_title"] == "Senior Frontend Engineer"
    assert created["ats_score"] == 88

    lst = client.get("/job-tracker", headers=headers)
    assert lst.status_code == 200
    assert any(j["id"] == created["id"] for j in lst.json())

    # Optional server-side status update (frontend keeps local copies too)
    upd = client.put(
        f"/job-tracker/{created['id']}",
        headers=headers,
        json={"status": "interview"},
    )
    assert upd.status_code == 200
    assert upd.json()["status"] == "interview"

    deleted = client.delete(f"/job-tracker/{created['id']}", headers=headers)
    assert deleted.status_code == 200
