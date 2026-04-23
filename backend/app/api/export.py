from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from app.models.schemas import ExportRequest
from app.services.pdf_export import create_resume_pdf
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

router = APIRouter()


def apply_changes_to_text(text: str, changes: list[dict]) -> str:
    """Apply changes to resume text."""
    result = text

    for change in changes:
        original = change.get("original", "")
        replacement = change.get("replacement", "")

        if original and replacement:
            # Simple replacement (case-insensitive for better matching)
            pattern = re.compile(re.escape(original), re.IGNORECASE)
            result = pattern.sub(replacement, result, count=1)

    return result


def create_docx_from_text(text: str) -> bytes:
    """Create a DOCX file from text."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Split text into sections
    lines = text.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check if line is a section header (all caps or starts with common headers)
        if line.isupper() or any(
            line.lower().startswith(header)
            for header in [
                "experience",
                "education",
                "skills",
                "summary",
                "objective",
                "projects",
                "certifications",
                "awards",
                "publications",
            ]
        ):
            # Add as heading
            p = doc.add_heading(line, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("-") or line.startswith("•") or line.startswith("*"):
            # Add as bullet point
            p = doc.add_paragraph(line[1:].strip(), style="List Bullet")
        else:
            # Add as regular paragraph
            p = doc.add_paragraph(line)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


@router.post("/export")
async def export_resume(request: ExportRequest):
    """
    Export resume as DOCX with applied changes.
    """
    # Validate input
    if not request.resume_text or not request.resume_text.strip():
        raise HTTPException(status_code=400, detail="Resume text is required")

    try:
        # Apply changes to text
        modified_text = apply_changes_to_text(
            request.resume_text, request.applied_changes
        )

        # Create DOCX
        docx_bytes = create_docx_from_text(modified_text)

        # Return as streaming response
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=optimized_resume.docx"
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exporting resume: {str(e)}")


@router.post("/export/pdf")
async def export_resume_pdf(request: ExportRequest):
    """
    Export resume as PDF with applied changes.
    """
    if not request.resume_text or not request.resume_text.strip():
        raise HTTPException(status_code=400, detail="Resume text is required")

    try:
        # Apply changes to text
        modified_text = apply_changes_to_text(
            request.resume_text, request.applied_changes
        )

        # Parse text into structured data for PDF
        # Simple parsing - extract sections from text
        lines = modified_text.split("\n")
        contact = {"name": "Resume"}
        summary = None
        skills = []
        experience = []
        education = []
        certifications = []

        current_section = None
        current_exp = None
        current_edu = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detect sections
            upper_line = line.upper()
            if "EXPERIENCE" in upper_line and len(line) < 30:
                current_section = "experience"
                continue
            elif "EDUCATION" in upper_line and len(line) < 30:
                current_section = "education"
                continue
            elif "SKILLS" in upper_line and len(line) < 30:
                current_section = "skills"
                continue
            elif "SUMMARY" in upper_line or "OBJECTIVE" in upper_line:
                current_section = "summary"
                continue
            elif "CERTIFICATION" in upper_line and len(line) < 30:
                current_section = "certifications"
                continue

            # Parse based on section
            if current_section == "summary":
                summary = line if not summary else summary + " " + line
            elif current_section == "skills":
                if line.startswith("-") or line.startswith("•"):
                    skills.append(line[1:].strip())
                else:
                    skills.extend([s.strip() for s in line.split(",") if s.strip()])
            elif current_section == "experience":
                if line.startswith("-") or line.startswith("•"):
                    if current_exp:
                        current_exp.setdefault("bullets", []).append(line[1:].strip())
                elif not current_exp:
                    current_exp = {"title": line}
                elif "company" not in current_exp:
                    current_exp["company"] = line
                else:
                    if "location" not in current_exp:
                        current_exp["location"] = line
                    elif "start_date" not in current_exp:
                        current_exp["start_date"] = line
                    elif "end_date" not in current_exp:
                        current_exp["end_date"] = line
                        experience.append(current_exp)
                        current_exp = None
            elif current_section == "education":
                if not current_edu:
                    current_edu = {"degree": line}
                elif "school" not in current_edu:
                    current_edu["school"] = line
                elif "location" not in current_edu:
                    current_edu["location"] = line
                elif "graduation_date" not in current_edu:
                    current_edu["graduation_date"] = line
                    education.append(current_edu)
                    current_edu = None
            elif current_section == "certifications":
                if line.startswith("-") or line.startswith("•"):
                    certifications.append(line[1:].strip())
                else:
                    certifications.append(line)

        # Create PDF
        pdf_bytes = create_resume_pdf(
            contact=contact,
            summary=summary,
            skills=skills if skills else None,
            experience=experience if experience else None,
            education=education if education else None,
            certifications=certifications if certifications else None,
        )

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": "attachment; filename=optimized_resume.pdf"
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exporting PDF: {str(e)}")
