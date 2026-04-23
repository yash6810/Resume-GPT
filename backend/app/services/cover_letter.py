import os
import json
from typing import Optional
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Load environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
HUGGINGFACE_API_TOKEN = os.getenv("HUGGINGFACE_API_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")


def generate_cover_letter_with_groq(
    resume_text: str, job_description: str, company_name: str, position: str
) -> Optional[str]:
    """Use Groq to generate a cover letter (fastest, unlimited free tier)."""
    if not GROQ_API_KEY:
        return None

    prompt = f"""Write a professional cover letter based on this resume and job description.

Resume:
{resume_text[:2000]}

Job Description:
{job_description[:2000]}

Company: {company_name}
Position: {position}

Write a compelling cover letter that:
1. Opens with a strong hook
2. Highlights relevant experience from the resume
3. Shows enthusiasm for the company
4. Closes with a call to action

Keep it concise (3-4 paragraphs)."""

    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "llama-3.1-8b-instant",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 800,
            },
            timeout=30,
        )

        if response.status_code == 200:
            result = response.json()
            return result["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"Groq error: {e}")

    return None


def generate_cover_letter_with_gemini(
    resume_text: str, job_description: str, company_name: str, position: str
) -> Optional[str]:
    """Use Gemini to generate a cover letter."""
    if not GEMINI_API_KEY:
        return None

    prompt = f"""Write a professional cover letter based on this resume and job description.

Resume:
{resume_text[:2000]}

Job Description:
{job_description[:2000]}

Company: {company_name}
Position: {position}

Write a compelling cover letter that:
1. Opens with a strong hook
2. Highlights relevant experience from the resume
3. Shows enthusiasm for the company
4. Closes with a call to action

Keep it concise (3-4 paragraphs)."""

    # Try multiple models in case of quota limits
    models = ["gemini-2.5-flash", "gemini-2.0-flash-lite", "gemini-2.0-flash"]

    for model in models:
        try:
            response = requests.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_API_KEY}",
                headers={"Content-Type": "application/json"},
                json={
                    "contents": [{"parts": [{"text": prompt}]}],
                    "generationConfig": {
                        "temperature": 0.7,
                        "maxOutputTokens": 800,
                    },
                },
                timeout=60,
            )

            if response.status_code == 200:
                result = response.json()
                return result["candidates"][0]["content"]["parts"][0]["text"]
            elif response.status_code == 429:
                # Quota exceeded, try next model
                continue
        except Exception as e:
            print(f"Gemini error ({model}): {e}")
            continue

    return None


def generate_cover_letter_with_openai(
    resume_text: str, job_description: str, company_name: str, position: str
) -> Optional[str]:
    """Use OpenAI to generate a cover letter."""
    if not OPENAI_API_KEY:
        return None

    prompt = f"""Write a professional cover letter based on this resume and job description.

Resume:
{resume_text[:2000]}

Job Description:
{job_description[:2000]}

Company: {company_name}
Position: {position}

Write a compelling cover letter that:
1. Opens with a strong hook
2. Highlights relevant experience from the resume
3. Shows enthusiasm for the company
4. Closes with a call to action

Keep it concise (3-4 paragraphs)."""

    try:
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {OPENAI_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "gpt-3.5-turbo",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 800,
            },
            timeout=60,
        )

        if response.status_code == 200:
            result = response.json()
            return result["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"OpenAI error: {e}")

    return None


def generate_cover_letter_with_huggingface(
    resume_text: str, job_description: str, company_name: str, position: str
) -> Optional[str]:
    """Use HuggingFace to generate a cover letter."""
    if not HUGGINGFACE_API_TOKEN:
        return None

    prompt = f"""Write a professional cover letter for {position} at {company_name}.

Resume highlights: {resume_text[:500]}

Job requirements: {job_description[:500]}

Cover letter:"""

    try:
        response = requests.post(
            "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.2",
            headers={
                "Authorization": f"Bearer {HUGGINGFACE_API_TOKEN}",
                "Content-Type": "application/json",
            },
            json={
                "inputs": prompt,
                "parameters": {
                    "max_new_tokens": 600,
                    "temperature": 0.7,
                    "return_full_text": False,
                },
            },
            timeout=60,
        )

        if response.status_code == 200:
            result = response.json()
            return result[0]["generated_text"].strip()
    except Exception as e:
        print(f"HuggingFace error: {e}")

    return None


def generate_template_cover_letter(
    resume_text: str, job_description: str, company_name: str, position: str
) -> str:
    """Generate a template-based cover letter when no LLM is available."""
    # Extract some keywords from job description
    common_skills = [
        "Python",
        "JavaScript",
        "React",
        "SQL",
        "AWS",
        "Docker",
        "Leadership",
        "Communication",
        "Problem-solving",
    ]

    matched_skills = []
    for skill in common_skills:
        if (
            skill.lower() in job_description.lower()
            or skill.lower() in resume_text.lower()
        ):
            matched_skills.append(skill)

    skills_text = (
        ", ".join(matched_skills[:3])
        if matched_skills
        else "relevant technical and soft skills"
    )

    return f"""Dear Hiring Manager,

I am writing to express my strong interest in the {position} position at {company_name}. With my background and experience, I am confident that I would be a valuable addition to your team.

Throughout my career, I have developed expertise in {skills_text}. My experience aligns well with the requirements outlined in your job description, and I am excited about the opportunity to contribute to {company_name}'s continued success.

I am particularly drawn to this role because it offers the opportunity to apply my skills in a dynamic environment. I am eager to bring my passion for excellence and collaborative spirit to your team.

I would welcome the opportunity to discuss how my background and skills would be a good fit for this position. Thank you for considering my application.

Sincerely,
[Your Name]"""


def generate_cover_letter(
    resume_text: str, job_description: str, company_name: str, position: str
) -> str:
    """Generate a cover letter using available methods.
    Priority: Groq (fastest) → Gemini → OpenAI → HuggingFace → Template
    """
    # Try Groq first (fastest, unlimited free tier)
    result = generate_cover_letter_with_groq(
        resume_text, job_description, company_name, position
    )
    if result:
        return result

    # Try Gemini only if Groq failed
    result = generate_cover_letter_with_gemini(
        resume_text, job_description, company_name, position
    )
    if result:
        return result

    # Try OpenAI
    result = generate_cover_letter_with_openai(
        resume_text, job_description, company_name, position
    )
    if result:
        return result

    # Try HuggingFace
    result = generate_cover_letter_with_huggingface(
        resume_text, job_description, company_name, position
    )
    if result:
        return result

    # Fall back to template
    return generate_template_cover_letter(
        resume_text, job_description, company_name, position
    )


def create_cover_letter_docx(cover_letter_text: str, applicant_name: str) -> bytes:
    """Create a DOCX file from cover letter text."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Split into paragraphs
    paragraphs = cover_letter_text.strip().split("\n\n")

    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.15

            # Handle the greeting and closing specially
            if para_text.strip().startswith("Dear") or para_text.strip().startswith(
                "Sincerely"
            ):
                run = para.add_run(para_text.strip())
            else:
                run = para.add_run(para_text.strip())

            run.font.size = Pt(11)
            run.font.name = "Calibri"

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
