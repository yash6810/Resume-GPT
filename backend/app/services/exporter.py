from docx import Document
from io import BytesIO

def generate_docx(resume_data: dict) -> BytesIO:
    """
    Generates a DOCX file from structured resume data.
    """
    document = Document()

    # Add a title
    document.add_heading('Resume', 0)

    # Add full text (optional, or can be used for sections not caught)
    # if 'resume_text' in resume_data and resume_data['resume_text'].strip():
    #     document.add_heading('Full Text', level=1)
    #     document.add_paragraph(resume_data['resume_text'])

    # Add sections
    if 'resume_sections' in resume_data:
        for section_name, section_content in resume_data['resume_sections'].items():
            if section_content.strip():
                # Capitalize the first letter of the section name for the heading
                document.add_heading(section_name.replace('_', ' ').title(), level=1)
                
                # Split content by lines to simulate bullet points or paragraphs
                lines = section_content.split('\n')
                for line in lines:
                    if line.strip(): # Only add non-empty lines
                        document.add_paragraph(line.strip(), style='List Bullet') # Use bullet style for readability
    
    # Save document to a BytesIO object
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0) # Rewind the stream to the beginning
    
    return file_stream
