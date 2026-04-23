from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from typing import List, Dict, Optional, Any


# Template color schemes
TEMPLATE_COLORS = {
    "modern": {
        "primary": RGBColor(44, 62, 80),  # Dark blue
        "secondary": RGBColor(52, 152, 219),  # Bright blue
        "accent": RGBColor(231, 76, 60),  # Red
        "text": RGBColor(44, 62, 80),
        "light": RGBColor(236, 240, 241),  # Light gray
    },
    "classic": {
        "primary": RGBColor(0, 0, 0),  # Black
        "secondary": RGBColor(51, 51, 51),  # Dark gray
        "accent": RGBColor(102, 102, 102),  # Medium gray
        "text": RGBColor(0, 0, 0),
        "light": RGBColor(245, 245, 245),  # Very light gray
    },
    "creative": {
        "primary": RGBColor(142, 68, 173),  # Purple
        "secondary": RGBColor(155, 89, 182),  # Light purple
        "accent": RGBColor(230, 126, 34),  # Orange
        "text": RGBColor(44, 62, 80),
        "light": RGBColor(245, 238, 248),  # Light purple
    },
    "minimal": {
        "primary": RGBColor(33, 33, 33),  # Almost black
        "secondary": RGBColor(97, 97, 97),  # Gray
        "accent": RGBColor(0, 150, 136),  # Teal
        "text": RGBColor(33, 33, 33),
        "light": RGBColor(250, 250, 250),  # Off-white
    },
    "executive": {
        "primary": RGBColor(0, 51, 102),  # Navy blue
        "secondary": RGBColor(70, 130, 180),  # Steel blue
        "accent": RGBColor(178, 34, 34),  # Firebrick red
        "text": RGBColor(33, 33, 33),
        "light": RGBColor(245, 245, 245),
    },
    "tech": {
        "primary": RGBColor(0, 200, 150),  # Cyan-green
        "secondary": RGBColor(100, 100, 100),  # Gray
        "accent": RGBColor(255, 99, 71),  # Tomato
        "text": RGBColor(33, 33, 33),
        "light": RGBColor(240, 255, 250),  # Honeydew
    },
    "academic": {
        "primary": RGBColor(139, 0, 0),  # Dark red
        "secondary": RGBColor(105, 105, 105),  # Dim gray
        "accent": RGBColor(0, 100, 0),  # Dark green
        "text": RGBColor(33, 33, 33),
        "light": RGBColor(255, 250, 240),  # Floral white
    },
}


def set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_horizontal_line(doc, color=None):
    """Add a horizontal line paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    # Default to black if no color provided
    if color is None:
        color_hex = "000000"
    else:
        # RGBColor stores as a 3-byte value
        color_hex = str(color)
    pBdr = parse_xml(
        f"<w:pBdr {nsdecls('w')}>"
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color_hex}"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def build_modern_template(data: Any) -> Document:
    """Modern template with colored header and clean sections."""
    doc = Document()
    colors = TEMPLATE_COLORS["modern"]

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Name - Large, bold, colored
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(4)
    run = name_para.add_run(data.contact.name.upper())
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = colors["primary"]

    # Contact info
    contact_parts = []
    if data.contact.email:
        contact_parts.append(data.contact.email)
    if data.contact.phone:
        contact_parts.append(data.contact.phone)
    if data.contact.location:
        contact_parts.append(data.contact.location)
    if data.contact.linkedin:
        contact_parts.append(data.contact.linkedin)
    if data.contact.website:
        contact_parts.append(data.contact.website)

    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_after = Pt(8)
    run = contact_para.add_run(" | ".join(contact_parts))
    run.font.size = Pt(10)
    run.font.color.rgb = colors["secondary"]

    # Horizontal line
    add_horizontal_line(doc, colors["secondary"])

    # Summary
    if data.summary:
        summary_heading = doc.add_paragraph()
        summary_heading.paragraph_format.space_before = Pt(12)
        summary_heading.paragraph_format.space_after = Pt(6)
        run = summary_heading.add_run("PROFESSIONAL SUMMARY")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        summary_para = doc.add_paragraph(data.summary)
        summary_para.paragraph_format.space_after = Pt(8)

    # Skills
    if data.skills:
        skills_heading = doc.add_paragraph()
        skills_heading.paragraph_format.space_before = Pt(12)
        skills_heading.paragraph_format.space_after = Pt(6)
        run = skills_heading.add_run("SKILLS")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.space_after = Pt(8)
        run = skills_para.add_run(" • ".join(data.skills))
        run.font.size = Pt(10)

    # Experience
    if data.experience:
        exp_heading = doc.add_paragraph()
        exp_heading.paragraph_format.space_before = Pt(12)
        exp_heading.paragraph_format.space_after = Pt(6)
        run = exp_heading.add_run("EXPERIENCE")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        for exp in data.experience:
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(8)
            exp_para.paragraph_format.space_after = Pt(2)

            title_run = exp_para.add_run(exp.title)
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = colors["primary"]

            exp_para.add_run("\n")
            company_run = exp_para.add_run(f"{exp.company} | {exp.location}")
            company_run.font.size = Pt(10)
            company_run.font.color.rgb = colors["secondary"]

            exp_para.add_run("\n")
            date_run = exp_para.add_run(f"{exp.start_date} - {exp.end_date}")
            date_run.font.size = Pt(10)
            date_run.italic = True

            for bullet in exp.bullets:
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(2)

    # Education
    if data.education:
        edu_heading = doc.add_paragraph()
        edu_heading.paragraph_format.space_before = Pt(12)
        edu_heading.paragraph_format.space_after = Pt(6)
        run = edu_heading.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        for edu in data.education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(6)
            edu_para.paragraph_format.space_after = Pt(2)

            degree_run = edu_para.add_run(edu.degree)
            degree_run.bold = True
            degree_run.font.size = Pt(11)

            edu_para.add_run("\n")
            school_run = edu_para.add_run(f"{edu.school} | {edu.location}")
            school_run.font.size = Pt(10)
            school_run.font.color.rgb = colors["secondary"]

            edu_para.add_run("\n")
            date_text = f"Graduated: {edu.graduation_date}"
            if edu.gpa:
                date_text += f" | GPA: {edu.gpa}"
            date_run = edu_para.add_run(date_text)
            date_run.font.size = Pt(10)
            date_run.italic = True

    # Certifications
    if data.certifications:
        cert_heading = doc.add_paragraph()
        cert_heading.paragraph_format.space_before = Pt(12)
        cert_heading.paragraph_format.space_after = Pt(6)
        run = cert_heading.add_run("CERTIFICATIONS")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        for cert in data.certifications:
            cert_para = doc.add_paragraph(cert, style="List Bullet")
            cert_para.paragraph_format.space_after = Pt(2)

    return doc


def build_classic_template(data: Any) -> Document:
    """Classic template with traditional formatting and horizontal lines."""
    doc = Document()
    colors = TEMPLATE_COLORS["classic"]

    # Standard margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Name - Centered, bold
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(4)
    run = name_para.add_run(data.contact.name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = colors["primary"]

    # Contact - Centered
    contact_parts = []
    if data.contact.email:
        contact_parts.append(data.contact.email)
    if data.contact.phone:
        contact_parts.append(data.contact.phone)
    if data.contact.location:
        contact_parts.append(data.contact.location)
    if data.contact.linkedin:
        contact_parts.append(data.contact.linkedin)
    if data.contact.website:
        contact_parts.append(data.contact.website)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(6)
    run = contact_para.add_run(" | ".join(contact_parts))
    run.font.size = Pt(10)

    # Horizontal line
    add_horizontal_line(doc, colors["primary"])

    # Summary
    if data.summary:
        heading = doc.add_heading("Professional Summary", level=2)
        heading.runs[0].font.color.rgb = colors["primary"]
        doc.add_paragraph(data.summary)

    # Skills
    if data.skills:
        heading = doc.add_heading("Skills", level=2)
        heading.runs[0].font.color.rgb = colors["primary"]
        skills_para = doc.add_paragraph()
        skills_para.add_run(", ".join(data.skills))

    # Experience
    if data.experience:
        heading = doc.add_heading("Experience", level=2)
        heading.runs[0].font.color.rgb = colors["primary"]

        for exp in data.experience:
            exp_para = doc.add_paragraph()
            title_run = exp_para.add_run(exp.title)
            title_run.bold = True

            exp_para.add_run(f"\n{exp.company} | {exp.location}")
            exp_para.add_run(f"\n{exp.start_date} - {exp.end_date}")

            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    # Education
    if data.education:
        heading = doc.add_heading("Education", level=2)
        heading.runs[0].font.color.rgb = colors["primary"]

        for edu in data.education:
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(edu.degree)
            degree_run.bold = True

            edu_para.add_run(f"\n{edu.school} | {edu.location}")
            date_text = f"\nGraduated: {edu.graduation_date}"
            if edu.gpa:
                date_text += f" | GPA: {edu.gpa}"
            edu_para.add_run(date_text)

    # Certifications
    if data.certifications:
        heading = doc.add_heading("Certifications", level=2)
        heading.runs[0].font.color.rgb = colors["primary"]
        for cert in data.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    return doc


def build_creative_template(data: Any) -> Document:
    """Creative template with accent colors and unique styling."""
    doc = Document()
    colors = TEMPLATE_COLORS["creative"]

    # Moderate margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Name - Large, creative font style
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    run = name_para.add_run(data.contact.name.upper())
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = colors["primary"]

    # Tagline/Title if available (using summary first line)
    if data.summary:
        tagline = data.summary.split(".")[0]
        tagline_para = doc.add_paragraph()
        tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline_para.paragraph_format.space_after = Pt(8)
        run = tagline_para.add_run(tagline)
        run.italic = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["secondary"]

    # Contact in a creative layout
    contact_parts = []
    if data.contact.email:
        contact_parts.append(f"✉ {data.contact.email}")
    if data.contact.phone:
        contact_parts.append(f"☎ {data.contact.phone}")
    if data.contact.location:
        contact_parts.append(f"📍 {data.contact.location}")
    if data.contact.linkedin:
        contact_parts.append(f"🔗 {data.contact.linkedin}")
    if data.contact.website:
        contact_parts.append(f"🌐 {data.contact.website}")

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(12)
    run = contact_para.add_run("  •  ".join(contact_parts))
    run.font.size = Pt(9)
    run.font.color.rgb = colors["accent"]

    # Decorative line
    add_horizontal_line(doc, colors["accent"])

    # Summary (full)
    if data.summary and len(data.summary) > len(data.summary.split(".")[0]):
        summary_heading = doc.add_paragraph()
        summary_heading.paragraph_format.space_before = Pt(10)
        summary_heading.paragraph_format.space_after = Pt(6)
        run = summary_heading.add_run("✦ ABOUT ME")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = colors["primary"]

        # Get full summary after tagline
        full_summary = ".".join(data.summary.split(".")[1:]).strip()
        if full_summary:
            if full_summary.startswith("."):
                full_summary = full_summary[1:]
            doc.add_paragraph(full_summary)

    # Skills with creative formatting
    if data.skills:
        skills_heading = doc.add_paragraph()
        skills_heading.paragraph_format.space_before = Pt(10)
        skills_heading.paragraph_format.space_after = Pt(6)
        run = skills_heading.add_run("✦ EXPERTISE")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = colors["primary"]

        # Skills in a flowing paragraph
        skills_para = doc.add_paragraph()
        for i, skill in enumerate(data.skills):
            if i > 0:
                skills_para.add_run(" • ")
            run = skills_para.add_run(skill)
            run.font.size = Pt(10)

    # Experience
    if data.experience:
        exp_heading = doc.add_paragraph()
        exp_heading.paragraph_format.space_before = Pt(10)
        exp_heading.paragraph_format.space_after = Pt(6)
        run = exp_heading.add_run("✦ EXPERIENCE")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = colors["primary"]

        for exp in data.experience:
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(10)
            exp_para.paragraph_format.space_after = Pt(2)

            title_run = exp_para.add_run(exp.title)
            title_run.bold = True
            title_run.font.size = Pt(12)
            title_run.font.color.rgb = colors["primary"]

            exp_para.add_run("\n")
            company_run = exp_para.add_run(f"{exp.company}")
            company_run.bold = True
            company_run.font.size = Pt(11)
            company_run.font.color.rgb = colors["secondary"]

            exp_para.add_run(f"  |  {exp.location}")
            exp_para.add_run(f"\n{exp.start_date} — {exp.end_date}")

            for bullet in exp.bullets:
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(2)

    # Education
    if data.education:
        edu_heading = doc.add_paragraph()
        edu_heading.paragraph_format.space_before = Pt(10)
        edu_heading.paragraph_format.space_after = Pt(6)
        run = edu_heading.add_run("✦ EDUCATION")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = colors["primary"]

        for edu in data.education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(6)

            degree_run = edu_para.add_run(edu.degree)
            degree_run.bold = True
            degree_run.font.size = Pt(11)

            edu_para.add_run(f"\n{edu.school}  |  {edu.location}")
            date_text = f"\n🎓 {edu.graduation_date}"
            if edu.gpa:
                date_text += f"  |  GPA: {edu.gpa}"
            edu_para.add_run(date_text)

    # Certifications
    if data.certifications:
        cert_heading = doc.add_paragraph()
        cert_heading.paragraph_format.space_before = Pt(10)
        cert_heading.paragraph_format.space_after = Pt(6)
        run = cert_heading.add_run("✦ CERTIFICATIONS")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = colors["primary"]

        for cert in data.certifications:
            doc.add_paragraph(f"🏅 {cert}")

    return doc


def build_minimal_template(data: Any) -> Document:
    """Minimal template with clean, sparse design."""
    doc = Document()
    colors = TEMPLATE_COLORS["minimal"]

    # Generous margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Name - Clean, simple
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(8)
    run = name_para.add_run(data.contact.name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = colors["primary"]

    # Contact - Minimal style
    contact_parts = []
    if data.contact.email:
        contact_parts.append(data.contact.email)
    if data.contact.phone:
        contact_parts.append(data.contact.phone)
    if data.contact.location:
        contact_parts.append(data.contact.location)
    if data.contact.linkedin:
        contact_parts.append(data.contact.linkedin)
    if data.contact.website:
        contact_parts.append(data.contact.website)

    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_after = Pt(16)
    run = contact_para.add_run(" · ".join(contact_parts))
    run.font.size = Pt(10)
    run.font.color.rgb = colors["secondary"]

    # Summary
    if data.summary:
        summary_para = doc.add_paragraph(data.summary)
        summary_para.paragraph_format.space_after = Pt(16)

    # Skills - Inline style
    if data.skills:
        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.space_after = Pt(16)
        run = skills_para.add_run(" · ".join(data.skills))
        run.font.size = Pt(10)
        run.font.color.rgb = colors["secondary"]

    # Experience
    if data.experience:
        for exp in data.experience:
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(10)
            exp_para.paragraph_format.space_after = Pt(2)

            # Title and company on same line
            title_run = exp_para.add_run(exp.title)
            title_run.bold = True
            title_run.font.size = Pt(11)

            exp_para.add_run(f" at {exp.company}")

            # Date and location
            date_para = doc.add_paragraph()
            date_para.paragraph_format.space_after = Pt(4)
            run = date_para.add_run(
                f"{exp.start_date} – {exp.end_date} · {exp.location}"
            )
            run.font.size = Pt(9)
            run.font.color.rgb = colors["secondary"]

            for bullet in exp.bullets:
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(2)

    # Education
    if data.education:
        edu_heading = doc.add_paragraph()
        edu_heading.paragraph_format.space_before = Pt(12)
        run = edu_heading.add_run("Education")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        for edu in data.education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_after = Pt(2)

            degree_run = edu_para.add_run(edu.degree)
            degree_run.bold = True
            degree_run.font.size = Pt(11)

            school_para = doc.add_paragraph()
            school_para.paragraph_format.space_after = Pt(2)
            run = school_para.add_run(f"{edu.school} · {edu.location}")
            run.font.size = Pt(10)
            run.font.color.rgb = colors["secondary"]

            date_text = edu.graduation_date
            if edu.gpa:
                date_text += f" · GPA: {edu.gpa}"
            date_para = doc.add_paragraph()
            date_para.paragraph_format.space_after = Pt(8)
            run = date_para.add_run(date_text)
            run.font.size = Pt(9)
            run.font.color.rgb = colors["secondary"]

    # Certifications
    if data.certifications:
        cert_heading = doc.add_paragraph()
        cert_heading.paragraph_format.space_before = Pt(12)
        run = cert_heading.add_run("Certifications")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        for cert in data.certifications:
            cert_para = doc.add_paragraph(cert)
            cert_para.paragraph_format.space_after = Pt(2)

    return doc


def build_executive_template(data: Any) -> Document:
    """Executive template with professional corporate styling."""
    doc = Document()
    colors = TEMPLATE_COLORS["executive"]

    # Narrow margins for more content
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Name - Bold, professional
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(2)
    run = name_para.add_run(data.contact.name.upper())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = colors["primary"]

    # Contact info
    contact_parts = []
    if data.contact.email:
        contact_parts.append(data.contact.email)
    if data.contact.phone:
        contact_parts.append(data.contact.phone)
    if data.contact.location:
        contact_parts.append(data.contact.location)

    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_after = Pt(6)
    run = contact_para.add_run(" | ".join(contact_parts))
    run.font.size = Pt(10)
    run.font.color.rgb = colors["secondary"]

    # LinkedIn and website on separate line
    links_parts = []
    if data.contact.linkedin:
        links_parts.append(data.contact.linkedin)
    if data.contact.website:
        links_parts.append(data.contact.website)
    if links_parts:
        links_para = doc.add_paragraph()
        links_para.paragraph_format.space_after = Pt(8)
        run = links_para.add_run(" | ".join(links_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = colors["secondary"]

    # Thick horizontal line
    add_horizontal_line(doc, colors["primary"])

    # Professional Summary
    if data.summary:
        summary_heading = doc.add_paragraph()
        summary_heading.paragraph_format.space_before = Pt(10)
        summary_heading.paragraph_format.space_after = Pt(6)
        run = summary_heading.add_run("PROFESSIONAL SUMMARY")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        summary_para = doc.add_paragraph(data.summary)
        summary_para.paragraph_format.space_after = Pt(8)

    # Core Competencies (Skills)
    if data.skills:
        skills_heading = doc.add_paragraph()
        skills_heading.paragraph_format.space_before = Pt(10)
        skills_heading.paragraph_format.space_after = Pt(6)
        run = skills_heading.add_run("CORE COMPETENCIES")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        # Skills in columns
        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.space_after = Pt(8)
        run = skills_para.add_run(" • ".join(data.skills))
        run.font.size = Pt(10)

    # Professional Experience
    if data.experience:
        exp_heading = doc.add_paragraph()
        exp_heading.paragraph_format.space_before = Pt(10)
        exp_heading.paragraph_format.space_after = Pt(6)
        run = exp_heading.add_run("PROFESSIONAL EXPERIENCE")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        for exp in data.experience:
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(8)
            exp_para.paragraph_format.space_after = Pt(2)

            title_run = exp_para.add_run(exp.title)
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = colors["primary"]

            exp_para.add_run(f"\n{exp.company} | {exp.location}")
            exp_para.add_run(f"\n{exp.start_date} - {exp.end_date}")

            for bullet in exp.bullets:
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(2)

    # Education
    if data.education:
        edu_heading = doc.add_paragraph()
        edu_heading.paragraph_format.space_before = Pt(10)
        edu_heading.paragraph_format.space_after = Pt(6)
        run = edu_heading.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        for edu in data.education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(6)
            edu_para.paragraph_format.space_after = Pt(2)

            degree_run = edu_para.add_run(edu.degree)
            degree_run.bold = True
            degree_run.font.size = Pt(11)

            edu_para.add_run(f"\n{edu.school} | {edu.location}")
            date_text = f"\nGraduated: {edu.graduation_date}"
            if edu.gpa:
                date_text += f" | GPA: {edu.gpa}"
            edu_para.add_run(date_text)

    # Certifications
    if data.certifications:
        cert_heading = doc.add_paragraph()
        cert_heading.paragraph_format.space_before = Pt(10)
        cert_heading.paragraph_format.space_after = Pt(6)
        run = cert_heading.add_run("CERTIFICATIONS")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        for cert in data.certifications:
            cert_para = doc.add_paragraph(cert, style="List Bullet")
            cert_para.paragraph_format.space_after = Pt(2)

    return doc


def build_tech_template(data: Any) -> Document:
    """Tech template with modern, clean design for tech professionals."""
    doc = Document()
    colors = TEMPLATE_COLORS["tech"]

    # Moderate margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Name - Clean, modern
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(2)
    run = name_para.add_run(data.contact.name)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = colors["primary"]

    # Contact info with icons
    contact_parts = []
    if data.contact.email:
        contact_parts.append(f"@ {data.contact.email}")
    if data.contact.phone:
        contact_parts.append(f"# {data.contact.phone}")
    if data.contact.location:
        contact_parts.append(f"$ {data.contact.location}")
    if data.contact.linkedin:
        contact_parts.append(f"& {data.contact.linkedin}")
    if data.contact.website:
        contact_parts.append(f"* {data.contact.website}")

    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_after = Pt(12)
    run = contact_para.add_run("  |  ".join(contact_parts))
    run.font.size = Pt(10)
    run.font.color.rgb = colors["secondary"]

    # Skills first (tech focus)
    if data.skills:
        skills_heading = doc.add_paragraph()
        skills_heading.paragraph_format.space_before = Pt(8)
        skills_heading.paragraph_format.space_after = Pt(6)
        run = skills_heading.add_run("TECH STACK")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        # Skills in a clean format
        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.space_after = Pt(12)
        for i, skill in enumerate(data.skills):
            if i > 0:
                skills_para.add_run("  •  ")
            run = skills_para.add_run(skill)
            run.font.size = Pt(10)
            run.font.color.rgb = colors["accent"]

    # Summary
    if data.summary:
        summary_heading = doc.add_paragraph()
        summary_heading.paragraph_format.space_before = Pt(8)
        summary_heading.paragraph_format.space_after = Pt(6)
        run = summary_heading.add_run("ABOUT")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        summary_para = doc.add_paragraph(data.summary)
        summary_para.paragraph_format.space_after = Pt(12)

    # Experience
    if data.experience:
        exp_heading = doc.add_paragraph()
        exp_heading.paragraph_format.space_before = Pt(8)
        exp_heading.paragraph_format.space_after = Pt(6)
        run = exp_heading.add_run("EXPERIENCE")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        for exp in data.experience:
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(8)
            exp_para.paragraph_format.space_after = Pt(2)

            title_run = exp_para.add_run(exp.title)
            title_run.bold = True
            title_run.font.size = Pt(11)

            exp_para.add_run(f"  @  {exp.company}")
            exp_para.add_run(f"\n{exp.start_date} - {exp.end_date}  |  {exp.location}")

            for bullet in exp.bullets:
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(2)

    # Education
    if data.education:
        edu_heading = doc.add_paragraph()
        edu_heading.paragraph_format.space_before = Pt(10)
        edu_heading.paragraph_format.space_after = Pt(6)
        run = edu_heading.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        for edu in data.education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(6)
            edu_para.paragraph_format.space_after = Pt(2)

            degree_run = edu_para.add_run(edu.degree)
            degree_run.bold = True
            degree_run.font.size = Pt(11)

            edu_para.add_run(f"\n{edu.school}  |  {edu.location}")
            date_text = f"\n{edu.graduation_date}"
            if edu.gpa:
                date_text += f"  |  GPA: {edu.gpa}"
            edu_para.add_run(date_text)

    # Certifications
    if data.certifications:
        cert_heading = doc.add_paragraph()
        cert_heading.paragraph_format.space_before = Pt(10)
        cert_heading.paragraph_format.space_after = Pt(6)
        run = cert_heading.add_run("CERTIFICATIONS")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = colors["primary"]

        for cert in data.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    return doc


def build_academic_template(data: Any) -> Document:
    """Academic CV template for research and academic positions."""
    doc = Document()
    colors = TEMPLATE_COLORS["academic"]

    # Standard margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Name - Centered, formal
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(4)
    run = name_para.add_run(data.contact.name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = colors["primary"]

    # Contact - Centered
    contact_parts = []
    if data.contact.email:
        contact_parts.append(data.contact.email)
    if data.contact.phone:
        contact_parts.append(data.contact.phone)
    if data.contact.location:
        contact_parts.append(data.contact.location)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(2)
    run = contact_para.add_run(" | ".join(contact_parts))
    run.font.size = Pt(10)

    # Links centered
    links_parts = []
    if data.contact.linkedin:
        links_parts.append(data.contact.linkedin)
    if data.contact.website:
        links_parts.append(data.contact.website)
    if links_parts:
        links_para = doc.add_paragraph()
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_para.paragraph_format.space_after = Pt(8)
        run = links_para.add_run(" | ".join(links_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = colors["secondary"]

    # Horizontal line
    add_horizontal_line(doc, colors["primary"])

    # Research Interests / Summary
    if data.summary:
        summary_heading = doc.add_paragraph()
        summary_heading.paragraph_format.space_before = Pt(10)
        summary_heading.paragraph_format.space_after = Pt(6)
        run = summary_heading.add_run("RESEARCH INTERESTS")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        summary_para = doc.add_paragraph(data.summary)
        summary_para.paragraph_format.space_after = Pt(8)

    # Education (important for academic CV)
    if data.education:
        edu_heading = doc.add_paragraph()
        edu_heading.paragraph_format.space_before = Pt(10)
        edu_heading.paragraph_format.space_after = Pt(6)
        run = edu_heading.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        for edu in data.education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(6)
            edu_para.paragraph_format.space_after = Pt(2)

            degree_run = edu_para.add_run(edu.degree)
            degree_run.bold = True
            degree_run.font.size = Pt(11)

            edu_para.add_run(f", {edu.school}")
            edu_para.add_run(f"\n{edu.location}")
            date_text = f"\n{edu.graduation_date}"
            if edu.gpa:
                date_text += f" | GPA: {edu.gpa}"
            edu_para.add_run(date_text)

    # Academic Experience
    if data.experience:
        exp_heading = doc.add_paragraph()
        exp_heading.paragraph_format.space_before = Pt(10)
        exp_heading.paragraph_format.space_after = Pt(6)
        run = exp_heading.add_run("ACADEMIC EXPERIENCE")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        for exp in data.experience:
            exp_para = doc.add_paragraph()
            exp_para.paragraph_format.space_before = Pt(8)
            exp_para.paragraph_format.space_after = Pt(2)

            title_run = exp_para.add_run(exp.title)
            title_run.bold = True
            title_run.font.size = Pt(11)

            exp_para.add_run(f"\n{exp.company}")
            exp_para.add_run(f"\n{exp.location}")
            exp_para.add_run(f"\n{exp.start_date} - {exp.end_date}")

            for bullet in exp.bullets:
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(2)

    # Technical Skills
    if data.skills:
        skills_heading = doc.add_paragraph()
        skills_heading.paragraph_format.space_before = Pt(10)
        skills_heading.paragraph_format.space_after = Pt(6)
        run = skills_heading.add_run("TECHNICAL SKILLS")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        skills_para = doc.add_paragraph()
        skills_para.paragraph_format.space_after = Pt(8)
        run = skills_para.add_run(", ".join(data.skills))
        run.font.size = Pt(10)

    # Certifications / Publications placeholder
    if data.certifications:
        cert_heading = doc.add_paragraph()
        cert_heading.paragraph_format.space_before = Pt(10)
        cert_heading.paragraph_format.space_after = Pt(6)
        run = cert_heading.add_run("CERTIFICATIONS & AWARDS")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = colors["primary"]

        for cert in data.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    return doc


def get_available_templates() -> List[Dict[str, str]]:
    """Return list of available templates with metadata."""
    return [
        {
            "id": "modern",
            "name": "Modern",
            "description": "Clean design with colored header and bullet-style skills",
        },
        {
            "id": "classic",
            "name": "Classic",
            "description": "Traditional centered header with horizontal lines",
        },
        {
            "id": "creative",
            "name": "Creative",
            "description": "Bold design with icons and accent colors",
        },
        {
            "id": "minimal",
            "name": "Minimal",
            "description": "Sparse, elegant layout with generous whitespace",
        },
        {
            "id": "executive",
            "name": "Executive",
            "description": "Professional corporate style with navy accents",
        },
        {
            "id": "tech",
            "name": "Tech",
            "description": "Modern tech-focused design with green accents",
        },
        {
            "id": "academic",
            "name": "Academic",
            "description": "Traditional academic CV style with dark red accents",
        },
    ]


def build_resume_with_template(data: Any) -> bytes:
    """Build resume DOCX using selected template."""
    template_id = data.template.lower()

    template_builders = {
        "modern": build_modern_template,
        "classic": build_classic_template,
        "creative": build_creative_template,
        "minimal": build_minimal_template,
        "executive": build_executive_template,
        "tech": build_tech_template,
        "academic": build_academic_template,
    }

    builder_func = template_builders.get(template_id, build_modern_template)
    doc = builder_func(data)

    import io

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# Industry Templates Support
def load_industry_templates() -> Dict[str, Any]:
    """Load industry templates from JSON file."""
    import json
    import os

    # Try multiple paths to find the industry_templates.json file
    possible_paths = [
        os.path.join(
            os.path.dirname(__file__), "..", "..", "data", "industry_templates.json"
        ),
        os.path.join(
            os.path.dirname(__file__),
            "..",
            "..",
            "..",
            "data",
            "industry_templates.json",
        ),
        os.path.join("data", "industry_templates.json"),
        os.path.join("..", "data", "industry_templates.json"),
    ]

    for path in possible_paths:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)

    # Return empty dict if file not found
    return {}


def get_industry_template(industry_id: str) -> Optional[Dict[str, Any]]:
    """Get a specific industry template by ID."""
    templates = load_industry_templates()
    return templates.get(industry_id)


def get_all_industry_templates() -> Dict[str, Any]:
    """Get all available industry templates."""
    return load_industry_templates()


def generate_industry_summary(
    industry_id: str, role: str = "", years: int = 0, skills: List[str] = None
) -> str:
    """Generate an industry-specific summary."""
    template = get_industry_template(industry_id)
    if not template:
        return ""

    summary_template = template.get("summary_template", "")
    if not summary_template:
        return ""

    # Fill in template placeholders
    summary = summary_template.replace("{role}", role or "professional")
    summary = summary.replace("{years}", str(years) if years > 0 else "5+")
    summary = summary.replace(
        "{skills}", ", ".join(skills[:3]) if skills else "relevant technologies"
    )
    summary = summary.replace("{achievement}", "delivering high-quality results")
    summary = summary.replace("{specialty}", template.get("name", "the industry"))
    summary = summary.replace("{skill}", skills[0] if skills else "key competencies")
    summary = summary.replace(
        "{interest}", "continuous learning and professional development"
    )
    summary = summary.replace("{setting}", "various settings")
    summary = summary.replace("{focus}", "student success")
    summary = summary.replace("{value}", "excellence")
    summary = summary.replace("{result}", "measurable outcomes")
    summary = summary.replace("{industries}", "multiple")

    return summary


def get_industry_keywords(industry_id: str) -> List[str]:
    """Get recommended keywords for an industry."""
    template = get_industry_template(industry_id)
    if not template:
        return []
    return template.get("keywords", [])


def get_industry_skills(industry_id: str) -> Dict[str, List[str]]:
    """Get recommended skills categories for an industry."""
    template = get_industry_template(industry_id)
    if not template:
        return {}
    return template.get("skills_categories", {})


def get_industry_bullet_templates(industry_id: str) -> List[str]:
    """Get bullet point templates for an industry."""
    template = get_industry_template(industry_id)
    if not template:
        return []
    return template.get("bullet_templates", [])
